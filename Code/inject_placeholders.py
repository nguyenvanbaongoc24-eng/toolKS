import docx
import os

def replace_text_in_docx(doc, replacements):
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                if p.text.strip() == old_text.strip():
                    p.text = p.text.replace(old_text, new_text)
                else:
                    for run in p.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            for run in p.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

def inject_phieu_khao_sat(out_dir):
    template_path = 'd:/ToolKS/Code/Phieu_Khao_Sat_ATTT_Mau_1.docx.docx'
    out_path = os.path.join(out_dir, 'phieu_khao_sat_template.docx')
    doc = docx.Document(template_path)
    
    replacements = {
        "ỦY BAN NHÂN DÂN XÃ LÝ NHÂN": "{{ ten_don_vi }}",
        "Ủy ban nhân dân xã Lý Nhân": "{{ ten_don_vi }}",
        "UBND xã Lý Nhân": "{{ ten_don_vi }}",
        "Thôn 5, xã Lý Nhân, tỉnh Ninh Bình": "{{ dia_chi }}",
        "ngày        tháng        năm 2026": "{{ ngay_khao_sat }}",
        "Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân": "{{ he_thong_thong_tin }}",
        "ông Đỗ Văn Hân": "{{ nguoi_khao_sat }}",
        "Trương Tuấn Lực": "{{ nguoi_dai_dien }}",
    }
    
    replace_text_in_docx(doc, replacements)
    doc.save(out_path)
    print(f"Created template: {out_path}")

def inject_hsdx(out_dir):
    template_path = 'd:/ToolKS/Code/HSDXCDAT LY NHAN.docx.docx'
    out_path = os.path.join(out_dir, 'hsdx_template.docx')
    doc = docx.Document(template_path)
    
    replacements = {
        "ỦY BAN NHÂN DÂN XÃ LÝ NHÂN": "{{ ten_don_vi }}",
        "UBND xã Lý Nhân": "{{ ten_don_vi }}",
        "xã Lý Nhân, tỉnh Ninh Bình": "{{ dia_chi }}",
        "Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân": "{{ he_thong_thong_tin }}",
        "ông Đỗ Văn Hân": "{{ nguoi_khao_sat }}",
        "Trương Tuấn Lực": "{{ nguoi_dai_dien }}",
        "năm 2026": "{{ nam_khao_sat }}",
    }
    
    # Add image placeholder for network diagram logic
    found = False
    for p in doc.paragraphs:
        if "4.1. Mô hình logic tổng thể" in p.text:
            p.add_run("\n{{ network_diagram }}")
            found = True
            break
    
    if not found:
        doc.add_paragraph("{{ network_diagram }}")

    replace_text_in_docx(doc, replacements)
    doc.save(out_path)
    print(f"Created template: {out_path}")

def inject_bao_cao(out_dir):
    template_path = 'd:/ToolKS/Code/BAO CAO HSDX.docx.docx'
    out_path = os.path.join(out_dir, 'bao_cao_template.docx')
    doc = docx.Document(template_path)
    
    replacements = {
        "Ủy ban nhân dân xã Lý Nhân": "{{ ten_don_vi }}",
        "UBND xã Lý Nhân": "{{ ten_don_vi }}",
        "ngày        tháng        năm 2026": "{{ ngay_khao_sat }}",
        "Thôn 5, xã Lý Nhân, tỉnh Ninh Bình": "{{ dia_chi }}",
    }
    
    replace_text_in_docx(doc, replacements)
    doc.save(out_path)
    print(f"Created template: {out_path}")

if __name__ == "__main__":
    out_dir = 'd:/ToolKS/Code/backend/templates'
    os.makedirs(out_dir, exist_ok=True)
    inject_phieu_khao_sat(out_dir)
    inject_hsdx(out_dir)
    inject_bao_cao(out_dir)
