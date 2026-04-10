from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import os
import logging
import traceback
from .diagram_generator import DiagramGenerator
from ai_extraction.security_analyzer import SecurityAnalyzer

logger = logging.getLogger(__name__)
# Avoid duplicate loggers if imported multiple times
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

class DocumentExporter:
    def __init__(self, template_dir="templates", output_dir="generated_docs"):
        # Make paths absolute relative to backend root
        backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.template_dir = os.path.join(backend_dir, "templates")
        self.output_dir = os.path.join(backend_dir, "generated_docs")
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)

    def _force_set_times_new_roman(self, doc):
        """
        Force Times New Roman font for all styles, paragraphs, and tables.
        """
        # 1. Set default document style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # 2. For Asian/Complex scripts support in Word
        rPr = font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # 3. Iterate through all paragraphs and set font explicitly for all runs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')

        # 4. Iterate through all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            rFonts.set(qn('w:cs'), 'Times New Roman')

    def _map_checkboxes(self, value, mapping):
        """
        Maps a form value to multiple Word placeholders for checkboxes.
        Uses 'V' (checkmark) instead of '☒'.
        """
        context = {}
        target_val = str(value or "").lower().strip()
        for option, placeholder in mapping.items():
            context[placeholder] = "V" if target_val == str(option).lower().strip() else "☐"
        return context

    def _get_hsdx_logic(self, data):
        """Logic for calculating Appendix I & II compliance in HSDX"""
        logic = {}
        
        # Appendix I & II Compliance Status - Updated logic based on new schema IDs
        logic['logic_policy_ok'] = "Đã đáp ứng" if data.get('k1_quy_che') and len(str(data.get('k1_quy_che', ''))) > 2 else "Chưa đáp ứng"
        logic['logic_vlan_ok'] = "Đã đáp ứng" if data.get('H4_co_vlan') == "Có" and data.get('T1_2_wifi_tach_rieng') == "Tách VLAN" else "Chưa đáp ứng"
        
        # Firewall check
        fw_type = data.get('l7_type') # Changed from L7_1_loai
        logic['logic_firewall_ok'] = "Đã đáp ứng" if fw_type and fw_type != "Không có" else "Chưa đáp ứng"
        
        logic['logic_av_ok'] = "Đã đáp ứng" if data.get('l3_av_has') == "Có" else "Chưa đáp ứng"
        logic['logic_pass_ok'] = "Đã đáp ứng" if data.get('L2_pass_policy') == "Có" else "Chưa đáp ứng"
        logic['logic_backup_ok'] = "Đã đáp ứng" if data.get('l4_bak_has') == "Tự động" else "Chưa đáp ứng"
        logic['logic_https_ok'] = "Đã đáp ứng" if data.get('p1_protocol') == "HTTPS (chứng chỉ SSL/TLS)" else "Chưa đáp ứng"
        
        logic['logic_training_ok'] = "Đã đáp ứng" if len(data.get('dao_tao') or []) >= 1 else "Chưa đáp ứng"
        logic['logic_audit_ok'] = "Đã đáp ứng" if len(data.get('kiem_tra_attt') or []) >= 1 else "Chưa đáp ứng"
        
        return logic

    def _get_report_logic(self, data):
        """Logic for generating issues and recommendations for the Report"""
        problems = []
        solutions = []

        # 1. Network Segmentation (Mục H4)
        if data.get('H4_co_vlan') != "Có":
            problems.append("- Hệ thống mạng chưa được thực hiện phân vùng mạng (VLAN) để cách ly các vùng chức năng.")
            solutions.append("- Thực hiện cấu hình phân chia VLAN trên các thiết bị Switch để tách biệt mạng máy chủ, máy trạm và chuyên dùng.")

        # 2. Wifi Separation (Mục T1.2)
        if data.get('T1_2_wifi_tach_rieng') == "Không" or data.get('T1_2_wifi_tach_rieng') == "Không có":
            pass 
        elif data.get('T1_2_wifi_tach_rieng') != "Tách VLAN":
            problems.append("- Mạng WiFi chưa được tách riêng hoàn toàn khỏi mạng LAN nội bộ nội bộ có dây của đơn vị.")
            solutions.append("- Quy hoạch lại mạng WiFi, sử dụng VLAN riêng và chính sách bảo mật WPA2/WPA3 Enterprise.")

        # 4. Firewall (Mục E2)
        if data.get('E2_firewall_type') == "Không có":
            problems.append("- Đơn vị chưa trang bị thiết bị Tường lửa (Firewall) chuyên dụng để kiểm soát lưu lượng.")
            solutions.append("- Trang bị thiết bị Tường lửa phần cứng (Next-Gen Firewall) để bảo vệ mạng biên đơn vị.")

        # 5. Antivirus (Mục L3)
        if data.get('l3_av_has') != "Có":
            problems.append("- Chưa trang bị phần mềm diệt virus tập trung trên các máy trạm tại đơn vị.")
            solutions.append("- Trang bị giải pháp diệt virus tập trung và thiết lập chế độ cập nhật định kỳ.")

        # 10. UPS (Mục L8.1)
        if data.get('L8_1_co_ups') != "Có":
            problems.append("- Phòng đặt thiết bị chưa có bộ lưu điện (UPS) đảm bảo duy trì nguồn điện khi có sự cố.")
            solutions.append("- Bổ sung thiết bị UPS có công suất phù hợp để bảo vệ phần cứng và duy trì mạng ổn định.")

        return {
            'problems': "\n".join(problems) if problems else "Hệ thống cơ bản ổn định.",
            'solutions': "\n".join(solutions) if solutions else "Tiếp tục duy trì vận hành ổn định.",
            'has_problems': len(problems) > 0
        }

    def _get_context(self, data):
        """
        Dynamically builds the context mapping based on input data with Smart Mapping.
        """
        # 1. Start with raw data
        context = {k: v for k, v in data.items() if not isinstance(v, (list, dict))}
        
        # 1.1 SMART MAPPING: Support both prefixed and unprefixed keys for Word templates
        # Example: 'A1_ten_don_vi' -> also provide 'ten_don_vi'
        items_to_add = {}
        for k, v in context.items():
            if '_' in k:
                parts = k.split('_', 1)
                prefix = parts[0]
                # Check if prefix is alphanumeric (like A1, L7, BC)
                if any(char.isdigit() for char in prefix) or len(prefix) <= 2:
                    plain_key = parts[1]
                    if plain_key not in context:
                        items_to_add[plain_key] = v
        context.update(items_to_add)

        # Apply defaults to ensure placeholders are replaced
        default_context = {
            'ten_don_vi': '...', 'dia_chi': '...', 'so_dien_thoai': '...', 'email': '...',
            'A1_ten_don_vi': '...', 'A2_ten_he_thong': '...', 'A3_dia_chi': '...',
            'nguoi_thuc_hien': '...', 'ngay_khao_sat': '...'
        }
        for k, v in default_context.items():
            if not context.get(k): context[k] = v

        # 2. Common Field Aliases
        context.update({
            'A1_ten_don_vi': data.get('A1_ten_don_vi') or data.get('ten_don_vi', '...'),
            'A2_ten_he_thong': data.get('A2_ten_he_thong') or data.get('he_thong_thong_tin', '...'),
            'A6_ho_ten_thu_truong': data.get('A6_ho_ten_thu_truong', '...'),
            'A6_chuc_vu_thu_truong': data.get('A6_chuc_vu_thu_truong', '...'),
            'nam_hien_tai': '2026',
        })

        # 3. Tables (FieldArrays)
        table_keys = {
            'can_bo_phu_trach': 'B_can_bo',
            'ket_noi_internet': 'D1_duong_truyen',
            'thiet_bi_mang': 'E1_thiet_bi_mang',
            'may_chu': 'F2_may_chu',
            'camera': 'G1_camera',
            'ip_tinh': 'H5_ip_tinh',
            'ung_dung': 'I1_ung_dung',
            'dao_tao': 'R1_dao_tao',
            'kiem_tra_attt': 'S1_kiem_tra',
            'port_switch': 'T2_port_mapping',
            'T5_vi_tri': 'T5_vi_tri'
        }
        for data_key, template_key in table_keys.items():
            rows = data.get(data_key, [])
            if isinstance(rows, list):
                context[template_key] = [{'idx': i+1, **row} for i, row in enumerate(rows)]
                # Also provide the data_key itself for older templates
                if data_key not in context:
                    context[data_key] = context[template_key]
            else:
                context[template_key] = []

        # 4. Photos (M1-M14) - Supporting both styles found in templates
        for i in range(1, 15):
            key = f"M{i}_status"
            is_present = data.get(key)
            context[f"{key}_ok"] = "V Đã có" if is_present else "☐ Chưa có"
            # Support more granular placeholders if they exist
            context[f"{key}_da_co"] = "V" if is_present else ""
            context[f"{key}_chua_co"] = "V" if not is_present else ""

        # 4.1 Section K Table (K_van_ban)
        k_mappings = [
            ('Quy chế ATTT', 'k1_quy_che'),
            ('Kế hoạch ATTT năm hiện tại', 'k2_ke_hoach_ht'),
            ('Kế hoạch ATTT năm trước', 'k3_ke_hoach_tr'),
            ('Quyết định phân công cán bộ', 'k4_qd_phan_cong_cb'),
            ('Quyết định phê duyệt hệ thống', 'K5_qd_phe_duyet_httt'),
            ('Quy trình ứng phó sự cố', 'K6_ung_pho_su_co'),
            ('Biền bản kiểm tra ATTT gần nhất', 'K7_bien_ban_kiem_tra'),
        ]
        context['K_van_ban'] = []
        for i, (label, key) in enumerate(k_mappings):
            val = data.get(key)
            context['K_van_ban'].append({
                'idx': i + 1,
                'loai': label,
                'so_vb': val or 'Chưa ban hành',
                'ghi_chu': 'Đã đính kèm' if val else 'Cần bổ sung'
            })

        # 5. Checkbox Logic - Mapping '☑' or '☐' based on schema_mappings.py
        # This makes the exporter truly dynamic and Python-native
        from .schema_mappings import CHECKBOX_MAPPINGS
        for field_id, mapping in CHECKBOX_MAPPINGS.items():
            val = data.get(field_id)
            for label, placeholder in mapping.items():
                is_checked = False
                if isinstance(val, list):
                    is_checked = label in val
                else:
                    is_checked = str(val) == str(label)
                context[placeholder] = "☑" if is_checked else "☐"

        # 6. Contact & Personnel Aliases
        context.update({
            'so_dien_thoai_co_quan': data.get('A4_so_dien_thoai') or '...',
            'email_co_quan': data.get('A5_email') or '...',
            'nguoi_khao_sat': data.get('N_nguoi_dien_ho_ten') or data.get('BC_nguoi_thuc_hien', '...'),
            'N_ngay_dien': data.get('N_ngay_dien') or data.get('BC_ngay_bao_cao', '...'),
            'ngay_lap_phieu': data.get('N_ngay_dien') or '...',
            'n_ngay_lap': data.get('BC_ngay_bao_cao') or data.get('N_ngay_dien') or '...',
        })

        # 7. Specialized Logic
        context.update(self._get_hsdx_logic(data))
        context.update(self._get_report_logic(data))

        return context

    def generate_phieu_khao_sat(self, data):
        template_path = os.path.join(self.template_dir, 'phieu_khao_sat_template.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
            
        doc = DocxTemplate(template_path)
        context = self._get_context(data)
        
        output_filename = f"Phieu_Khao_Sat_{data.get('ten_don_vi', 'NoName')}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        try:
            doc.render(context)
            self._force_set_times_new_roman(doc)
            doc.save(output_path)
            return output_path
        except Exception as e:
            logger.error(f"Render error: {str(e)}")
            raise e

    def generate_hsdx(self, data):
        template_path = os.path.join(self.template_dir, 'hsdx_template.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
            
        doc = DocxTemplate(template_path)
        context = self._get_context(data)
        
        # Generate Diagram logic remains same
        try:
            dg = DiagramGenerator(output_dir=self.output_dir)
            logic_path = dg.generate_logical_diagram(data, f"logic_{data.get('ten_don_vi', 'img')}.png")
            context['network_diagram'] = InlineImage(doc, logic_path, width=Mm(150))
        except:
            context['network_diagram'] = ""
        
        output_filename = f"HSDX_{data.get('ten_don_vi', 'NoName')}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        doc.render(context)
        self._force_set_times_new_roman(doc)
        doc.save(output_path)
        return output_path

    def generate_bao_cao(self, data):
        template_path = os.path.join(self.template_dir, 'bao_cao_template.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
            
        doc = DocxTemplate(template_path)
        context = self._get_context(data)
        
        output_filename = f"Bao_Cao_{data.get('ten_don_vi', 'NoName')}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        doc.render(context)
        self._force_set_times_new_roman(doc)
        doc.save(output_path)
        return output_path
