import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


def generate_survey_report(data: dict, output_path: str) -> str:
    """
    Generate a BAO CAO HSDX-style report from survey form data.
    Matches the structure of the MobiFone survey report template.
    """
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    ten_don_vi = data.get('ten_don_vi', 'ỦY BAN NHÂN DÂN')
    he_thong = data.get('he_thong_thong_tin', 'Hệ thống thông tin')
    dia_chi = data.get('dia_chi', '')
    nguoi_dung_dau = data.get('nguoi_dung_dau', '')
    sdt = data.get('so_dien_thoai', '')
    email = data.get('email', '')

    # ===== HEADER =====
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_left.add_run("ĐƠN VỊ TƯ VẤN ATTT")
    run.font.size = Pt(12)
    run.bold = True

    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc")
    run.font.size = Pt(12)
    run.bold = True

    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            tc = cell._tc
            tcPr = tc.tcPr if tc.tcPr is not None else tc._new_tcPr()
            from docx.oxml.ns import qn
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is not None:
                tcPr.remove(borders)

    doc.add_paragraph("")

    # Subject line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"V/v báo cáo khảo sát thực trạng, nhu cầu chuyển đổi số và đánh giá an toàn thông tin cho {ten_don_vi}")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(f"Kính gửi: {ten_don_vi}")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("")

    # ===== I. CĂN CỨ PHÁP LÝ =====
    p = doc.add_paragraph()
    run = p.add_run("I. CĂN CỨ PHÁP LÝ")
    run.bold = True
    run.font.size = Pt(13)

    legal_bases = [
        "Căn cứ Luật An toàn thông tin mạng ngày 19 tháng 11 năm 2015;",
        "Căn cứ Luật An ninh mạng ngày 12/6/2018;",
        "Căn cứ Nghị định số 85/2016/NĐ-CP ngày 01/7/2016 của Chính phủ về bảo đảm an toàn hệ thống thông tin theo cấp độ;",
        "Căn cứ Thông tư số 12/2022/TT-BTTTT ngày 12/8/2022 của Bộ Thông tin và Truyền thông;",
    ]
    for lb in legal_bases:
        p = doc.add_paragraph(f"- {lb}")
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")

    # ===== II. MỤC ĐÍCH =====
    p = doc.add_paragraph()
    run = p.add_run("II. MỤC ĐÍCH, YÊU CẦU, PHẠM VI KHẢO SÁT")
    run.bold = True

    purposes = [
        "Đánh giá hiện trạng bảo đảm an toàn thông tin của hệ thống thông tin đang được quản lý, vận hành.",
        "Xác định mức độ đáp ứng các yêu cầu về an toàn thông tin theo cấp độ quy định.",
        "Phát hiện các nguy cơ, lỗ hổng, rủi ro mất an toàn thông tin trong quá trình vận hành.",
        "Đề xuất các biện pháp, giải pháp kỹ thuật, quản lý nhằm nâng cao mức độ bảo đảm ATTT.",
    ]
    for pur in purposes:
        p = doc.add_paragraph(f"- {pur}")
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")

    # ===== III. PHƯƠNG PHÁP =====
    p = doc.add_paragraph()
    run = p.add_run("III. PHƯƠNG PHÁP KHẢO SÁT")
    run.bold = True

    p = doc.add_paragraph("Gồm 03 phương pháp: Khảo sát thực tế tại hiện trường; Phỏng vấn cán bộ phụ trách CNTT; Kiểm tra cấu hình, thông số kỹ thuật hệ thống.")

    doc.add_paragraph("")

    # ===== IV. KẾT QUẢ KHẢO SÁT =====
    p = doc.add_paragraph()
    run = p.add_run("IV. KẾT QUẢ KHẢO SÁT HIỆN TRẠNG HẠ TẦNG CNTT")
    run.bold = True

    # --- Thông tin đơn vị ---
    p = doc.add_paragraph()
    run = p.add_run("1. Thông tin đơn vị")
    run.bold = True
    
    info_items = [
        ("Tên đơn vị", ten_don_vi),
        ("Địa chỉ", dia_chi),
        ("Người đứng đầu", nguoi_dung_dau),
        ("Điện thoại", sdt),
        ("Email", email),
        ("Hệ thống thông tin", he_thong),
    ]
    for label, val in info_items:
        if val:
            doc.add_paragraph(f"- {label}: {val}")

    doc.add_paragraph("")

    # --- Bảng thiết bị mạng ---
    thiet_bi = data.get('thiet_bi_mang', [])
    if thiet_bi:
        p = doc.add_paragraph()
        run = p.add_run("2. Danh sách trang thiết bị CNTT")
        run.bold = True

        cols = ['STT', 'Loại thiết bị', 'Hãng SX', 'Model', 'Số Serial', 'Vị trí lắp đặt', 'Năm mua', 'Ghi chú']
        table = doc.add_table(rows=1 + len(thiet_bi), cols=len(cols))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for ci, col_name in enumerate(cols):
            cell = table.cell(0, ci)
            cell.text = col_name
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Rows
        for ri, tb in enumerate(thiet_bi):
            vals = [
                str(ri + 1),
                tb.get('loai_thiet_bi', ''),
                tb.get('hang', ''),
                tb.get('model', ''),
                tb.get('serial', ''),
                tb.get('vi_tri', ''),
                tb.get('nam_mua', ''),
                tb.get('ghi_chu', ''),
            ]
            for ci, v in enumerate(vals):
                cell = table.cell(ri + 1, ci)
                cell.text = v
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph("")

    # --- Bảng máy chủ ---
    may_chu = data.get('may_chu', [])
    if may_chu:
        p = doc.add_paragraph()
        run = p.add_run("3. Danh sách máy chủ")
        run.bold = True

        cols = ['STT', 'Vai trò', 'Hãng / Model', 'Số Serial', 'RAM (GB)', 'Ổ cứng (TB)', 'Hệ điều hành', 'Vị trí']
        table = doc.add_table(rows=1 + len(may_chu), cols=len(cols))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ci, col_name in enumerate(cols):
            cell = table.cell(0, ci)
            cell.text = col_name
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for ri, mc in enumerate(may_chu):
            vals = [
                str(ri + 1),
                mc.get('vai_tro', ''),
                f"{mc.get('hang', '')} {mc.get('model', '')}".strip(),
                mc.get('serial', ''),
                mc.get('ram', ''),
                mc.get('hdd', ''),
                mc.get('he_dieu_hanh', ''),
                mc.get('vi_tri', ''),
            ]
            for ci, v in enumerate(vals):
                cell = table.cell(ri + 1, ci)
                cell.text = v
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph("")

    # --- Bảng Camera ---
    cameras = data.get('camera', [])
    if cameras:
        p = doc.add_paragraph()
        run = p.add_run("4. Danh sách Camera")
        run.bold = True

        cols = ['STT', 'Hãng SX', 'Model', 'Số Serial', 'Độ phân giải', 'Vị trí lắp đặt', 'Ghi chú']
        table = doc.add_table(rows=1 + len(cameras), cols=len(cols))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ci, col_name in enumerate(cols):
            cell = table.cell(0, ci)
            cell.text = col_name
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for ri, cam in enumerate(cameras):
            vals = [
                str(ri + 1),
                cam.get('hang', ''),
                cam.get('model', ''),
                cam.get('serial', ''),
                cam.get('do_phan_giai', ''),
                cam.get('vi_tri', ''),
                cam.get('ghi_chu', ''),
            ]
            for ci, v in enumerate(vals):
                cell = table.cell(ri + 1, ci)
                cell.text = v
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph("")

    # --- Hiện trạng bảo mật ---
    p = doc.add_paragraph()
    run = p.add_run("5. Hiện trạng bảo mật")
    run.bold = True

    security_items = [
        ("Tường lửa (Firewall)", data.get('tuong_lua_loai', 'Chưa rõ')),
        ("Phần mềm Anti-virus", data.get('anti_virus', 'Chưa rõ')),
        ("Chính sách mật khẩu mạnh", data.get('chinh_sach_mat_khau', 'Chưa rõ')),
        ("Mã hóa dữ liệu", data.get('ma_hoa_du_lieu', 'Chưa rõ')),
        ("VPN", data.get('vpn', 'Chưa rõ')),
        ("Sao lưu dữ liệu", data.get('sao_luu', 'Chưa rõ')),
        ("Cập nhật HĐH", data.get('cap_nhat_he_dieu_hanh', 'Chưa rõ')),
    ]
    for label, val in security_items:
        doc.add_paragraph(f"- {label}: {val}")

    doc.add_paragraph("")

    # ===== V. ĐÁNH GIÁ CHUNG =====
    p = doc.add_paragraph()
    run = p.add_run("V. ĐÁNH GIÁ CHUNG")
    run.bold = True

    p = doc.add_paragraph(f"Sau khi khảo sát và đánh giá trên kết quả khảo sát thì hệ thống của {ten_don_vi} đang còn tồn tại các vấn đề cần khắc phục để đáp ứng yêu cầu cấp độ an toàn thông tin theo quy định.")

    # Auto-generate issues based on data
    issues = []
    if not data.get('tuong_lua_loai') or data.get('tuong_lua_loai', '').strip() == '':
        issues.append("Chưa có thiết bị tường lửa (Firewall) chuyên dụng để kiểm soát truy cập giữa các vùng mạng.")
    if data.get('anti_virus') == 'Không' or data.get('anti_virus') == 'Có - Miễn phí':
        issues.append("Các máy tính người dùng chưa được cài đặt phần mềm diệt virus bản quyền.")
    if data.get('chinh_sach_mat_khau') == 'Không':
        issues.append("Chưa có chính sách mật khẩu mạnh thống nhất.")
    if data.get('sao_luu') == 'Không':
        issues.append("Chưa có quy trình sao lưu dữ liệu định kỳ.")
    if data.get('vpn') == 'Không':
        issues.append("Chưa triển khai VPN cho kết nối từ xa.")
    if not issues:
        issues.append("Hệ thống cơ sở hạ tầng cần được nâng cấp thêm để đáp ứng yêu cầu cấp độ 2.")

    for issue in issues:
        doc.add_paragraph(f"+ {issue}")

    p = doc.add_paragraph()
    run = p.add_run(f"Kết luận chung: Hệ thống cơ sở hạ tầng CNTT của {ten_don_vi} cần được bổ sung, nâng cấp để đáp ứng yêu cầu về an toàn thông tin theo cấp độ.")
    run.bold = True

    doc.add_paragraph("")

    # ===== VI. ĐỀ XUẤT =====
    p = doc.add_paragraph()
    run = p.add_run("VI. ĐỀ XUẤT")
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("1. Danh sách các trang thiết bị cần trang bị")
    run.bold = True

    # Auto-generate investment proposals
    proposals = [
        ["1", "Firewall", "01", "Kiểm soát truy cập Internet; Firewall SPI; VPN", "Cao"],
        ["2", "Switch Layer 2 có hỗ trợ VLAN", "06", "Phân tách vùng mạng: LAN, WiFi, Camera", "Cao"],
        ["3", "Phần mềm diệt virus bản quyền", "01 gói", "Phòng chống mã độc trên toàn bộ PC", "Cao"],
        ["4", "Ổ cứng di động 1TB (sao lưu)", "02", "Sao lưu định kỳ dữ liệu, luân phiên", "Cao"],
        ["5", "UPS (Bộ lưu điện dự phòng)", "01", "Bảo vệ thiết bị mạng khỏi mất điện đột ngột", "Cao"],
        ["6", "Rack server", "01", "Lắp đặt tập trung thiết bị mạng", "Cao"],
    ]

    proposal_cols = ['STT', 'Tên thiết bị/Giải pháp', 'Số lượng', 'Mục đích đầu tư', 'Ưu tiên']
    table = doc.add_table(rows=1 + len(proposals), cols=len(proposal_cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, col_name in enumerate(proposal_cols):
        cell = table.cell(0, ci)
        cell.text = col_name
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for ri, row_data in enumerate(proposals):
        for ci, v in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = v
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")

    # ===== BẢNG CAM KẾT =====
    p = doc.add_paragraph()
    run = p.add_run("2. Nội dung cam kết thực hiện")
    run.bold = True

    commitments = [
        ["1", "Ban hành Quy chế bảo đảm an toàn, an ninh thông tin mạng", "Trong vòng 06 tháng", ten_don_vi],
        ["2", "Ban hành Quyết định phân công cán bộ phụ trách CNTT/ATTT", "Trong vòng 06 tháng", ten_don_vi],
        ["3", "Xây dựng quy trình sao lưu dữ liệu và quy trình ứng cứu sự cố", "Trong vòng 06 tháng", ten_don_vi],
        ["4", "Tổ chức tập huấn/phổ biến ATTT cho các cán bộ", "Trong vòng 06 tháng", ten_don_vi],
        ["5", "Mua sắm và triển khai phần mềm diệt virus", "Trong vòng 12 tháng", ten_don_vi],
        ["6", "Mua sắm Router/Firewall chuyên dụng", "Trong vòng 18 tháng", ten_don_vi],
        ["7", "Triển khai phân vùng VLAN", "Trong vòng 18 tháng", ten_don_vi],
        ["8", "Mua sắm ổ cứng dự phòng và thực hiện sao lưu định kỳ", "Trong vòng 06 tháng", ten_don_vi],
    ]

    commit_cols = ['STT', 'Nội dung cam kết', 'Thời hạn thực hiện', 'Đơn vị chịu trách nhiệm']
    table = doc.add_table(rows=1 + len(commitments), cols=len(commit_cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, col_name in enumerate(commit_cols):
        cell = table.cell(0, ci)
        cell.text = col_name
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for ri, row_data in enumerate(commitments):
        for ci, v in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = v
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")

    # ===== CLOSING =====
    p = doc.add_paragraph(f"Trên đây là báo cáo khảo sát thực trạng hệ thống ATTT của {ten_don_vi}./.")
    doc.add_paragraph("")
    p = doc.add_paragraph("Trân trọng!")

    doc.add_paragraph("")

    # Footer table
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    left = footer_table.cell(0, 0)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Nơi nhận:\n- Như trên;\n- Lưu: VT.")
    run.font.size = Pt(10)
    run.italic = True

    right = footer_table.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NGƯỜI THỰC HIỆN BÁO CÁO")
    run.bold = True
    run.font.size = Pt(13)

    # Remove footer table borders
    for row in footer_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.tcPr if tc.tcPr is not None else tc._new_tcPr()

    doc.save(output_path)
    return output_path
