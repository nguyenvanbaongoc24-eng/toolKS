import os
from docx import Document

def create_survey_template():
    # Setup paths
    base_dir = os.path.dirname(os.path.abspath(__file__))
    templates_dir = os.path.join(base_dir, 'templates')
    os.makedirs(templates_dir, exist_ok=True)
    template_path = os.path.join(templates_dir, 'default_template.docx')
    
    doc = Document()
    
    # Title
    doc.add_heading('PHIẾU KHẢO SÁT AN TOÀN THÔNG TIN', 0)
    
    # Section A
    doc.add_heading('A. THÔNG TIN CHUNG', level=1)
    doc.add_paragraph('Tên đơn vị: {{ ten_don_vi }}')
    doc.add_paragraph('Người đứng đầu: {{ nguoi_dung_dau }}')
    doc.add_paragraph('Địa chỉ: {{ dia_chi }}')
    doc.add_paragraph('Điện thoại: {{ so_dien_thoai }}')
    doc.add_paragraph('Email: {{ email }}')
    doc.add_paragraph('Tổng số cán bộ: {{ so_can_bo }}')
    
    # Section B
    doc.add_heading('B. KẾT NỐI INTERNET', level=1)
    doc.add_paragraph('Nhà cung cấp: {{ ket_noi_internet.nha_cung_cap }}')
    doc.add_paragraph('Loại đường truyền: {{ ket_noi_internet.loai_ket_noi }}')
    doc.add_paragraph('Băng thông: {{ ket_noi_internet.bang_thong }}')
    doc.add_paragraph('Địa chỉ IP WAN: {{ ket_noi_internet.ip_wan }}')

    # Group C & D & E for lists
    doc.add_heading('DANH SÁCH THIẾT BỊ', level=1)
    
    # Section E Servers
    doc.add_heading('Máy chủ', level=2)
    doc.add_paragraph('{% tr for mc in may_chu %}Tên: {{ mc.ten }} | OS: {{ mc.he_dieu_hanh }} | RAM: {{ mc.ram }} | HDD: {{ mc.o_cung }}{% endtr %}')

    # Section C Network
    doc.add_heading('Thiết bị mạng & An toàn bảo mật', level=2)
    doc.add_paragraph('{% tr for tb in thiet_bi_mang %}Loại: {{ tb.loai_thiet_bi }} | Hãng: {{ tb.hang_san_xuat }} | Model: {{ tb.model }} | Vị trí: {{ tb.vi_tri }}{% endtr %}')
    
    # Section D Camera
    doc.add_heading('Camera', level=2)
    doc.add_paragraph('{% tr for cam in camera %}Vị trí: {{ cam.vi_tri }} | Hãng: {{ cam.hang }} | Model: {{ cam.model }}{% endtr %}')
    
    # Section F Apps
    doc.add_heading('F. PHẦN MỀM & ỨNG DỤNG', level=1)
    doc.add_paragraph('{% tr for ung_dung in ung_dung %}Tên: {{ ung_dung.ten_ung_dung }} | CC: {{ ung_dung.don_vi_cung_cap }}{% endtr %}')

    # Final Sections
    doc.add_heading('G. NHÂN SỰ VÀ GHI CHÚ', level=1)
    doc.add_paragraph('Thông tin nhân sự ATTT: {{ nhan_su }}')
    doc.add_paragraph('Ghi chú thêm: {{ ghi_chu }}')
    
    doc.save(template_path)
    print(f"Created template successfully at: {template_path}")

if __name__ == "__main__":
    create_survey_template()
