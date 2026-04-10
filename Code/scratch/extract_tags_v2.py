import re
from docx import Document
import os

def find_tags_in_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return set()
    
    doc = Document(file_path)
    tags = set()
    
    # Improved regex
    pattern = r'\{\{.*?\}\}|\{\%.*?\%\}'
    
    def process_text(text):
        if not text: return
        matches = re.findall(pattern, text, re.DOTALL)
        for m in matches:
            tags.add(m.strip())

    # Check paragraphs
    for p in doc.paragraphs:
        process_text(p.text)
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text(cell.text)
                    
    return tags

templates = [
    'Code/backend/templates/phieu_khao_sat_template.docx',
]

for t in templates:
    print(f"\n--- Tags in {t} ---")
    ts = sorted(list(find_tags_in_docx(t)))
    for tag in ts:
        print(tag)
