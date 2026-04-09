from docx import Document
import os

path = os.path.join('Code', 'Mẫu phiếu khảo sát.docx')
doc = Document(path)
with open('Code/scratch/template_text.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text.strip() + '\n')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        f.write(f"[TABLE] {p.text.strip()}\n")
