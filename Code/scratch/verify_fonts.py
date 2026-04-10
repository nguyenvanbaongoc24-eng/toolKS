from docx import Document
import os

def check_fonts_in_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = Document(file_path)
    fonts = set()
    
    # Check paragraphs
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.name:
                            fonts.add(r.font.name)
                            
    print(f"Fonts found in {os.path.basename(file_path)}: {fonts}")

check_fonts_in_docx('Code/backend/generated_docs/Phieu_Khao_Sat_Test Don Vi.docx')
