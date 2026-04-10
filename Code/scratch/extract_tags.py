import re
from docx import Document
import os

def find_tags_in_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return set()
    
    doc = Document(file_path)
    tags = set()
    
    # regex to find {{ something }} and {% something %}
    pattern = r'(\{\{[^}]+\}\}|\{\%[^%]+\%\})'
    
    # Check paragraphs
    for p in doc.paragraphs:
        matches = re.findall(pattern, p.text)
        for m in matches:
            tags.add(m.strip())
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for m in matches:
                    tags.add(m.strip())
                    
    return tags

templates = [
    'Code/backend/templates/phieu_khao_sat_template.docx',
    'Code/backend/templates/hsdx_template.docx',
    'Code/backend/templates/bao_cao_template.docx'
]

for t in templates:
    print(f"\n--- Tags in {t} ---")
    ts = sorted(list(find_tags_in_docx(t)))
    for tag in ts:
        print(tag)
