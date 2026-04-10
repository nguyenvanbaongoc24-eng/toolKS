import re
from docx import Document
import os

def find_placeholders_in_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return set()
    
    doc = Document(file_path)
    placeholders = set()
    
    # regex to find {{ something }}
    pattern = r'\{\{([^}]+)\}\}'
    
    # Check paragraphs
    for p in doc.paragraphs:
        matches = re.findall(pattern, p.text)
        for m in matches:
            placeholders.add(m.strip())
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for m in matches:
                    placeholders.add(m.strip())
                    
    return placeholders

templates = [
    'Code/backend/templates/phieu_khao_sat_template.docx',
    'Code/backend/templates/hsdx_template.docx',
    'Code/backend/templates/bao_cao_template.docx'
]

for t in templates:
    print(f"\n--- Placeholders in {t} ---")
    ps = sorted(list(find_placeholders_in_docx(t)))
    for p in ps:
        if '.' in p: continue # Skip logic/filters for now
        print(p)
