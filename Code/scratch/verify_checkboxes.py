import re
from docx import Document
import os

def check_placeholders_in_docx(file_path, expected_checked):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = Document(file_path)
    found_checked = {}
    
    def process_text(text):
        if not text: return
        # Find any remaining {{ ... }}
        matches = re.findall(r'\{\{.*?\}\}', text)
        if matches:
            print(f"STILL HAS PLACEHOLDERS: {matches}")
            
        for key in expected_checked:
            if key in text:
                print(f"Found placeholder text {key} in: {text[:100]}")

    # For simplicity, we just search for the strings "☑" and "☐"
    checked_count = 0
    unchecked_count = 0
    
    for p in doc.paragraphs:
        checked_count += p.text.count("☑")
        unchecked_count += p.text.count("☐")
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                checked_count += cell.text.count("☑")
                unchecked_count += cell.text.count("☐")
                    
    print(f"Summary for {os.path.basename(file_path)}:")
    print(f"Checked (☑): {checked_count}")
    print(f"Unchecked (☐): {unchecked_count}")

# We expect at least a few checked boxes (E2_firewall_cung_co, L5_2_thoi_gian_3_6, L5_ghi_log_co, L6_1_su_co_khong)
check_placeholders_in_docx('Code/backend/generated_docs/Phieu_Khao_Sat_Test Don Vi.docx', ['L5_2_thoi_gian_3_6'])
