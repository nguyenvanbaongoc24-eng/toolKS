import os
import re
import logging
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

def get_full_text_from_runs(runs):
    return ''.join(run.text for run in runs)

def replace_text_preserving_format(paragraph, old_text, new_text):
    full_text = get_full_text_from_runs(paragraph.runs)
    if old_text not in full_text:
        return False
    new_full_text = full_text.replace(old_text, new_text)
    if not paragraph.runs: return False
    paragraph.runs[0].text = new_full_text
    for run in paragraph.runs[1:]: run.text = ''
    return True

def replace_in_document(doc, replacements):
    count = 0
    sorted_old_texts = sorted(replacements.keys(), key=len, reverse=True)
    for old_text in sorted_old_texts:
        new_text = replacements[old_text]
        for paragraph in doc.paragraphs:
            if old_text in get_full_text_from_runs(paragraph.runs):
                if replace_text_preserving_format(paragraph, old_text, new_text): count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in get_full_text_from_runs(paragraph.runs):
                            if replace_text_preserving_format(paragraph, old_text, new_text): count += 1
    return count

def inject_table_loop(doc, table_index, loop_var, row_placeholders):
    if table_index >= len(doc.tables): return False
    table = doc.tables[table_index]
    if len(table.rows) < 2: return False
    
    # 1. Identify the sample row (original Row 1)
    sample_row = table.rows[1]
    
    # 2. Insert helper row BEFORE for the 'start' tag
    # docxtpl {%tr ... %} will strip the entire row it's in
    start_row = table.add_row()
    # Move it to before sample_row
    sample_row._element.addprevious(start_row._element)
    start_row.cells[0].paragraphs[0].text = f"{{%tr for item in {loop_var} %}}"
    
    # 3. Insert helper row AFTER for the 'end' tag
    end_row = table.add_row()
    # Move it to after sample_row
    sample_row._element.addnext(end_row._element)
    end_row.cells[0].paragraphs[0].text = "{%tr endfor %}"
    
    # 4. Inject placeholders into the sample row (now Row 2 essentially)
    last_idx = min(len(sample_row.cells), len(row_placeholders)) - 1
    for i in range(last_idx + 1):
        cell = sample_row.cells[i]
        placeholder = row_placeholders[i]
        
        # Cleanup paragraphs
        for p_idx in range(len(cell.paragraphs) - 1, 0, -1):
            p_xml = cell.paragraphs[p_idx]._element
            p_xml.getparent().remove(p_xml)
            
        p = cell.paragraphs[0]
        p.text = placeholder
            
    logger.info(f"  Tagged Table {table_index} with Triple-Row loop: {loop_var}")
    return True

def get_replacements():
    r = {
        "A1. Tên đơn vị chủ quản hệ thống thông tin (*)  xã Lý Nhân, tỉnh Ninh Bình": "A1. Tên đơn vị chủ quản hệ thống thông tin (*)  {{ ten_don_vi }}",
        "A2. Tên hệ thống thông tin cần phân loại (*)  Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân": "A2. Tên hệ thống thông tin cần phân loại (*)  {{ he_thong_thong_tin }}",
        "A3. Địa chỉ trụ sở chính (*)  thôn 5, xã Lý Nhân, tỉnh Ninh Bình ": "A3. Địa chỉ trụ sở chính (*)  {{ dia_chi }}",
        " Số điện thoại cơ quan: không có": " Số điện thoại cơ quan: {{ so_dien_thoai_co_quan }}",
        " Email cơ quan: vanthu.lynhan.gov.ninhbinh.vn": " Email cơ quan: {{ email_co_quan }}",
        "Ông Trương Tuấn Lực – Chủ tịch Ủy Ban Nhân dân xã Lý Nhân": "{{ A6_ho_ten_thu_truong }} – {{ A6_chuc_vu_thu_truong }}",
        "Hệ thống phục vụ hoạt động hành chính của UBND xã: xử lý hồ sơ một cửa (DVCTT), quản lý văn bản điều hành, lưu trữ dữ liệu dân cư, trao đổi email nội bộ": "{{ C1_mo_ta_chuc_nang }}",
        "Cán bộ công chức xã (33 người), người dân đến giao dịch một cửa, UBND xã.": "{{ C2_doi_tuong_nguoi_dung }}",
        "Dữ liệu cá nhân (CMND, địa chỉ, số điện thoại); hồ sơ hộ tịch; hồ sơ đất đai; văn bản hành chính (không mật).": "{{ C3_loai_du_lieu }}",
        "Nội bộ: 33 cán bộ; Bên ngoài: ~200 lượt/tháng (người dân qua cổng DVCTT)": "Nội bộ: {{ C5_noi_bo }} cán bộ; Bên ngoài: {{ C5_ben_ngoai }}",
        " Năm 2025": " năm {{ C6_nam_hoat_dong }}",
        "Hệ thống DVCTT tỉnh": "{{ C7_ten_he_thong_cap_tren }}",
        "iGate 10.66.138.209": "{{ D2_router_modem }}",
        "10.66.138.210": "{{ D3_ip_lan_gateway }}",
        
        # Checkboxes Section L1
        "☐ Có khóa cửa (chìa khóa thường)": "{{ L1_khoa_cua_thuong }} Có khóa cửa (chìa khóa thường)",
        "☐ Có khóa cửa + camera giám sát": "{{ L1_khoa_camera }} Có khóa cửa + camera giám sát",
        "☐ Có thẻ từ / kiểm soát điện tử": "{{ L1_the_tu }} Có thẻ từ / kiểm soát điện tử",
        "☐ Không có kiểm soát riêng": "{{ L1_khong_kiem_soat }} Không có kiểm soát riêng",
        
        # Section L2
        "☐ Có – Yêu cầu: độ dài tối thiểu _____ ký tự, đổi định kỳ _____ tháng": "{{ L2_pass_policy_co }} Có – Yêu cầu: độ dài tối thiểu {{ L2_do_dai_min }} ký tự, đổi định kỳ {{ L2_doi_dinh_ky_thang }} tháng",
        "☐ Không có chính sách thống nhất (mỗi người tự đặt)": "{{ L2_pass_policy_khong }} Không có chính sách thống nhất (mỗi người tự đặt)",
        
        # Rack
        "Kích thước rack: ___U  Vị trí: Tầng ___ – Phòng ___": "Kích thước rack: {{ T3_1_rack_u }}U  Vị trí: Tầng {{ T3_1_rack_tang }} – Phòng {{ T3_1_rack_phong }}",
        
        # Dates & Footer
        "ngày        tháng        năm 2026": "ngày {{ N_ngay_dien }}",
        "Trương Tuấn Lực": "{{ A6_ho_ten_thu_truong }}",
        "Đỗ Văn Hân": "{{ nguoi_khao_sat }}", # User might be using BC_nguoi_thuc_hien too
    }
    # Loop over common checkbox glyphs
    for glyph in ["☐", "☒"]:
        r[f"{glyph} Cá nhân thông thường"] = "{{ C4_du_lieu_ca_nhan_thuong }} Cá nhân thông thường"
        r[f"{glyph} Dữ liệu cá nhân thông thường"] = "{{ C4_du_lieu_ca_nhan_thuong }} Dữ liệu cá nhân thông thường"
        r[f"{glyph} Dữ liệu cá nhân nhạy cảm"] = "{{ C4_du_lieu_ca_nhan_nhay_cam }} Dữ liệu cá nhân nhạy cảm"
        r[f"{glyph} Dữ liệu công"] = "{{ C4_du_lieu_cong }} Dữ liệu công"
        r[f"{glyph} Không xác định"] = "{{ C4_khong_xac_dinh }} Không xác định"
        r[f"{glyph} Có – Tên hệ thống:"] = "{{ C7_co }} Có – Tên hệ thống: {{ C7_ten_he_thong_cap_tren }}"
    return r

def inject_phieu(src_path, out_path):
    logger.info(f"=== Processing Phieu: {src_path} ===")
    doc = Document(src_path)
    replace_in_document(doc, get_replacements())
    # Corrected indices based on all_headers.txt
    inject_table_loop(doc, 3, "can_bo_phu_trach", ["{{item.idx}}", "{{item.ho_ten}}", "{{item.chuc_vu}}", "{{item.so_dt}}", "{{item.email}}", "{{item.trinh_do}}", "{{item.chung_chi_attt}}"])
    inject_table_loop(doc, 6, "ket_noi_internet", ["{{item.idx}}", "{{item.isp}}", "{{item.loai_ket_noi}}", "{{item.bang_thong}}", "{{item.vai_tro}}", "{{item.ip_wan}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 8, "thiet_bi_mang", ["{{item.idx}}", "{{item.loai_thiet_bi}}", "{{item.hang_san_xuat}}", "{{item.model}}", "{{item.so_serial}}", "{{item.vi_tri}}", "{{item.nam_mua}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 12, "may_chu", ["{{item.idx}}", "{{item.vai_tro}}", "{{item.hang_model}}", "{{item.so_serial}}", "{{item.ram_gb}}", "{{item.o_cung_tb}}", "{{item.he_dieu_hanh}}", "{{item.vi_tri}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 14, "camera", ["{{item.idx}}", "{{item.hang_san_xuat}}", "{{item.model}}", "{{item.so_serial}}", "{{item.do_phan_giai}}", "{{item.vi_tri}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 16, "ip_tinh", ["{{item.idx}}", "{{item.ten_thiet_bi}}", "{{item.dia_chi_ip}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 18, "ung_dung", ["{{item.idx}}", "{{item.ten}}", "{{item.chuc_nang}}", "{{item.don_vi_cung_cap}}", "{{item.phien_ban}}", "{{item.ket_noi_internet}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 20, "K_van_ban", ["{{item.idx}}", "{{item.loai}}", "{{item.so_vb}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 26, "dao_tao", ["{{item.idx}}", "{{item.hinh_thuc}}", "{{item.don_vi_to_chuc}}", "{{item.thoi_gian}}", "{{item.so_can_bo}}", "{{item.chung_chi_so_vb}}"])
    inject_table_loop(doc, 28, "kiem_tra_attt", ["{{item.idx}}", "{{item.loai_kiem_tra}}", "{{item.don_vi_thuc_hien}}", "{{item.thoi_gian}}", "{{item.ket_qua_so_vb}}"])
    inject_table_loop(doc, 31, "port_switch", ["{{item.ten_switch}}", "{{item.so_cong}}", "{{item.cong_su_dung}}", "{{item.ghi_chu}}"])
    inject_table_loop(doc, 33, "T5_vi_tri", ["{{item.idx}}", "{{item.thiet_bi}}", "{{item.tang}}", "{{item.phong}}", "{{item.ghi_chu}}"])
    doc.save(out_path)
    logger.info(f"  Saved Phieu template to: {out_path}")

if __name__ == "__main__":
    src_dir = "d:/ToolKS/Code"
    out_dir = "d:/ToolKS/Code/backend/templates"
    os.makedirs(out_dir, exist_ok=True)
    phieu_src = os.path.join(src_dir, "Phieu_Khao_Sat_ATTT_Mau_1.docx.docx")
    if os.path.exists(phieu_src): inject_phieu(phieu_src, os.path.join(out_dir, "phieu_khao_sat_template.docx"))
    logger.info("=== Template injection complete! ===")
