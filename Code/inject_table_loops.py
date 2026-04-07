from docx import Document
import os
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def inject_table_loops(doc_path):
    doc = Document(doc_path)
    count = 0
    
    # Define signatures for tables we want to loop over:
    # Tuple: (header_keyword, loop_tag, end_tag)
    # The script will look for header_keyword in the first row. 
    # If found, it injects loop_tag into the first cell of the second row, 
    # and end_tag into the last cell of the second row.
    
    signatures = [
        ('Nhà cung cấp (ISP)', '{% tr for d in D1_duong_truyen %}', '{% tr endfor %}'),
        ('Loại thiết bị', '{% tr for d in thiet_bi_mang %}', '{% tr endfor %}'),
        ('Hệ điều hành chính', '{% tr for d in may_tinh %}', '{% tr endfor %}'), # PC
        ('Vai trò (*) (File server', '{% tr for d in may_chu %}', '{% tr endfor %}'),
        ('Camera', '{% tr for c in camera %}', '{% tr endfor %}'), 
        ('Độ phân giải', '{% tr for c in camera %}', '{% tr endfor %}'),
        ('Địa chỉ IP tĩnh', '{% tr for ip in ip_tinh %}', '{% tr endfor %}'),
        ('Ứng dụng', '{% tr for u in ung_dung %}', '{% tr endfor %}'),
        ('Đơn vị tổ chức', '{% tr for d in dao_tao %}', '{% tr endfor %}'),
        ('Kết quả / Số văn bản kết luận', '{% tr for k in kiem_tra %}', '{% tr endfor %}'),
        ('Cổng đang sử dụng', '{% tr for p in T2_port_mapping %}', '{% tr endfor %}')
    ]
    
    for table in doc.tables:
        if len(table.rows) > 1:
            header_text = "".join([c.text for c in table.rows[0].cells])
            
            for sig, start_tag, end_tag in signatures:
                if sig in header_text:
                    row2 = table.rows[1]
                    first_cell = row2.cells[0]
                    last_cell = row2.cells[-1]
                    
                    # Prevent double injection
                    if "{% tr for" not in first_cell.text:
                        # Prepend start tag
                        if len(first_cell.paragraphs) > 0:
                            p = first_cell.paragraphs[0]
                            p.insert_paragraph_before(start_tag)
                        else:
                            first_cell.text = start_tag + "\n" + first_cell.text
                            
                        # Append end tag
                        if len(last_cell.paragraphs) > 0:
                            last_cell.add_paragraph(end_tag)
                        else:
                            last_cell.text = last_cell.text + "\n" + end_tag
                            
                        count += 1
                        logger.info(f"Injecting {start_tag} into table with header [{sig}]")
                        break
                        
    if count > 0:
        doc.save(doc_path)
        logger.info(f"Updated {count} tables in {os.path.basename(doc_path)}")
        
if __name__ == "__main__":
    templates_dir = r"d:\ToolKS\Code\backend\templates"
    for fname in ["phieu_khao_sat_template.docx", "hsdx_template.docx"]:
        path = os.path.join(templates_dir, fname)
        if os.path.exists(path):
            inject_table_loops(path)
