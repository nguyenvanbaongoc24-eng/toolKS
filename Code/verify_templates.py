"""Quick verification: Check if templates contain Jinja2 placeholders."""
from docx import Document
import os

templates = [
    ("phieu_khao_sat_template.docx", "Phiếu Khảo Sát"),
    ("hsdx_template.docx", "HSDX"),
    ("bao_cao_template.docx", "Báo Cáo")
]

template_dir = r"d:\ToolKS\Code\backend\templates"

for filename, label in templates:
    path = os.path.join(template_dir, filename)
    doc = Document(path)
    
    placeholders_found = []
    old_text_remaining = []
    
    for p in doc.paragraphs:
        # Check for Jinja2 placeholders
        if "{{" in p.text and "}}" in p.text:
            import re
            tags = re.findall(r'\{\{[^}]+\}\}', p.text)
            placeholders_found.extend(tags)
        
        # Check for old hardcoded text
        if "Lý Nhân" in p.text:
            old_text_remaining.append(p.text[:80])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text and "}}" in p.text:
                        import re
                        tags = re.findall(r'\{\{[^}]+\}\}', p.text)
                        placeholders_found.extend(tags)
                    if "Lý Nhân" in p.text:
                        old_text_remaining.append(f"[TABLE] {p.text[:80]}")
    
    print(f"\n{'='*60}")
    print(f"📄 {label} ({filename})")
    print(f"{'='*60}")
    
    # Deduplicate  
    unique_placeholders = sorted(set(placeholders_found))
    print(f"✅ Found {len(unique_placeholders)} unique placeholders:")
    for tag in unique_placeholders[:20]:
        print(f"   {tag}")
    if len(unique_placeholders) > 20:
        print(f"   ... and {len(unique_placeholders) - 20} more")
    
    if old_text_remaining:
        print(f"\n⚠️  Still has {len(old_text_remaining)} instances of 'Lý Nhân':")
        for text in old_text_remaining[:5]:
            print(f"   {text}")
    else:
        print(f"\n✅ No old 'Lý Nhân' text remaining!")
