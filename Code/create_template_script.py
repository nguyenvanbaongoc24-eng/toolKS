import os
import sys
from docxtpl import DocxTemplate
from docx import Document

# Mapping from sample strings to placeholders
PARAGRAPH_REPLACEMENTS = {
    "xã Lý Nhân, tỉnh Ninh Bình": "{{ ten_don_vi }}",
    "Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân": "{{ he_thong_thong_tin }}",
    "thôn 5, xã Lý Nhân, tỉnh Ninh Bình": "{{ dia_chi }}",
    "không có": "{{ so_dien_thoai }}",
    "vanthu.lynhan.gov.ninhbinh.vn": "{{ email }}",
    "Ông Trương Tuấn Lực – Chủ tịch Ủy Ban Nhân dân xã Lý Nhân": "{{ A6_ho_ten_thu_truong }}",
    "[Số QĐ, ngày ký, cơ quan ký]": "{{ A7_so_quyet_dinh }}",
    "iGate 10.66.138.209": "{{ D2_router_modem }}",
    "[10.66.138.210]": "{{ D3_ip_lan_gateway }}",
    "192.168.1.0/24": "{{ H1_dai_ip_lan }}",
    "192.168.1.1": "{{ H2_ip_gateway }}",
    "8.8.8.8 (Google)": "{{ H3_dns }}",
    "Họ tên: ___________________": "Họ tên: {{ n_nguoi_lap }}",
    "Chức vụ: __________________": "Chức vụ: {{ n_chuc_vu_lap }}",
}

def inject_table_tags(table, iter_name, field_mappings):
    """
    Injects {% for item in iter_name %} and {% endfor %} into a table.
    Assumes the first row is header, and data starts from the last empty-ish row.
    """
    if len(table.rows) < 2:
        return
    
    # We'll put the for loop in the second row (usually the first data row)
    first_data_row = table.rows[1]
    last_data_row = table.rows[-1]
    
    # Inject tags into cells
    for i, (col_idx, field_name) in enumerate(field_mappings.items()):
        if col_idx < len(first_data_row.cells):
            cell = first_data_row.cells[col_idx]
            cell.text = f"{{{{ item.{field_name} }}}}"
    
    # Add Jinja control tags
    # Actually docxtpl works better if tags are in the first/last cells of the row
    first_data_row.cells[0].text = f"{{% for item in {iter_name} %}}" + first_data_row.cells[0].text
    first_data_row.cells[-1].text = first_data_row.cells[-1].text + f"{{% endfor %}}"

def main():
    sample_path = "Phieu_Khao_Sat_ATTT_Mau_1.docx.docx"
    output_path = "backend/templates/phieu_khao_sat_template.docx"
    
    if not os.path.exists(sample_path):
        print(f"Sample not found: {sample_path}")
        return

    doc = Document(sample_path)

    # 1. Replace paragraph text
    for p in doc.paragraphs:
        for old, new in PARAGRAPH_REPLACEMENTS.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    # 2. Inject Table Tags
    # Table 4: Cán bộ (STT, Họ tên, Chức vụ, SĐT, Email, Trình độ, Chứng chỉ)
    # Based on phieu_ks_placeholders.txt
    inject_table_tags(doc.tables[3], "B_can_bo", {
        0: "idx", 1: "ho_ten", 2: "chuc_vu", 3: "so_dt", 4: "email", 5: "trinh_do", 6: "chung_chi_attt"
    })

    # Table 7: Đường truyền
    inject_table_tags(doc.tables[6], "D1_duong_truyen", {
        0: "idx", 1: "isp", 2: "loai_ket_noi", 3: "bang_thong", 4: "vai_tro", 5: "ip_wan", 6: "ghi_chu"
    })

    # Table 9: Thiết bị mạng
    inject_table_tags(doc.tables[8], "E1_thiet_bi_mang", {
        0: "idx", 1: "loai_thiet_bi", 2: "hang_san_xuat", 3: "model", 4: "so_serial", 5: "vi_tri", 6: "nam_mua", 7: "ghi_chu"
    })

    # Table 13: Máy chủ
    inject_table_tags(doc.tables[12], "F2_may_chu", {
        0: "idx", 1: "vai_tro", 2: "hang_model", 3: "so_serial", 4: "ram_gb", 5: "o_cung_tb", 6: "he_dieu_hanh", 7: "vi_tri"
    })

    # Table 15: Camera
    inject_table_tags(doc.tables[14], "G1_camera", {
        0: "idx", 1: "hang_san_xuat", 2: "model", 3: "so_serial", 4: "do_phan_giai", 5: "vi_tri", 6: "ghi_chu"
    })

    # Table 17: IP Tĩnh
    inject_table_tags(doc.tables[16], "H5_ip_tinh", {
        0: "idx", 1: "ten_thiet_bi", 2: "ip_tinh", 3: "ghi_chu"
    })

    # Table 19: Ứng dụng
    inject_table_tags(doc.tables[18], "I1_ung_dung", {
        0: "idx", 1: "ten_ung_dung", 2: "chuc_nang", 3: "nha_cung_cap", 4: "version", 5: "inter_conn", 6: "ghi_chu"
    })

    # Table 27: Đào tạo
    inject_table_tags(doc.tables[26], "R1_dao_tao", {
        0: "idx", 1: "hinh_thuc", 2: "to_chuc", 3: "thoi_gian", 4: "so_cb", 5: "chung_chi"
    })

    # Table 29: Kiểm tra
    inject_table_tags(doc.tables[28], "S1_kiem_tra", {
        0: "idx", 1: "loai_kt", 2: "thuc_hien", 3: "thoi_gian", 4: "ket_qua"
    })

    # Table 32: Port Mapping
    inject_table_tags(doc.tables[31], "T2_port_mapping", {
        0: "sw_name", 1: "port_count", 2: "port_desc", 3: "ghi_chu"
    })

    # Table 34: Vị trí thiết bị
    inject_table_tags(doc.tables[33], "T5_vi_tri", {
        0: "ten_thiet_bi", 1: "tang", 2: "phong", 3: "ghi_chu"
    })

    # 3. Handle Section M Checkboxes (Table 36)
    m_table = doc.tables[35]
    for i in range(1, 15):
        if i < len(m_table.rows):
            row = m_table.rows[i]
            # Column 3 is the status/checkbox area
            if len(row.cells) > 3:
                 row.cells[3].text = f"{{{{ M{i}_status_ok }}}}"

    doc.save(output_path)
    print(f"Template created: {output_path}")

if __name__ == "__main__":
    main()
