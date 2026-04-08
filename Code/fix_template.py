import os
from docx import Document

def clean_and_inject_triple_row(table, iter_name, field_mappings):
    """
    Proven Triple Row method for docxtpl row loops:
    1. Header (0)
    2. Row with {%tr for ... %} (index 1)
    3. Row with {{ item.fields }} (index 2)
    4. Row with {%tr endfor %} (index 3)
    """
    print(f"Applying Triple Row to {iter_name}...")
    
    # 1. Clear all rows except header
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    
    # 2. Add 3 rows
    r_start = table.add_row()
    r_data = table.add_row()
    r_end = table.add_row()
    
    # 3. Inject Start Loop {%tr for item in items %}
    # Must be {%tr WITHOUT SPACE for the preprocessor to work reliably
    r_start.cells[0].text = f"{{%tr for item in {iter_name} %}}"
    # Clear other cells in loop row just in case
    for i in range(1, len(r_start.cells)):
        r_start.cells[i].text = ""

    # 4. Inject Data Row {{ item.field }}
    for col_idx, field_name in field_mappings.items():
        if col_idx < len(r_data.cells):
            r_data.cells[col_idx].text = f"{{{{ item.{field_name} }}}}"

    # 5. Inject End Loop {%tr endfor %}
    r_end.cells[0].text = "{%tr endfor %}"
    for i in range(1, len(r_end.cells)):
        r_end.cells[i].text = ""

def main():
    sample_path = r"d:\ToolKS\Code\Phieu_Khao_Sat_ATTT_Mau_1.docx.docx"
    output_path = r"d:\ToolKS\Code\backend\templates\phieu_khao_sat_template.docx"
    
    if not os.path.exists(sample_path):
        print(f"Sample not found: {sample_path}")
        return

    doc = Document(sample_path)
    
    # Full list of tables to fix
    tables_to_fix = [
        (3, "B_can_bo", {0: "idx", 1: "ho_ten", 2: "chuc_vu", 3: "so_dt", 4: "email", 5: "trinh_do", 6: "chung_chi_attt"}),
        (6, "D1_duong_truyen", {0: "idx", 1: "isp", 2: "loai_ket_noi", 3: "bang_thong", 4: "vai_tro", 5: "ip_wan", 6: "ghi_chu"}),
        (8, "E1_thiet_bi_mang", {0: "idx", 1: "loai_thiet_bi", 2: "hang_san_xuat", 3: "model", 4: "so_serial", 5: "vi_tri", 6: "nam_mua", 7: "ghi_chu"}),
        (12, "F2_may_chu", {0: "idx", 1: "vai_tro", 2: "hang_model", 3: "so_serial", 4: "ram_gb", 5: "o_cung_tb", 6: "he_dieu_hanh", 7: "vi_tri"}),
        (14, "G1_camera", {0: "idx", 1: "hang_san_xuat", 2: "model", 3: "so_serial", 4: "do_phan_giai", 5: "vi_tri", 6: "ghi_chu"}),
        (16, "H5_ip_tinh", {0: "idx", 1: "ten_thiet_bi", 2: "ip_tinh", 3: "ghi_chu"}),
        (18, "I1_ung_dung", {0: "idx", 1: "ten_ung_dung", 2: "chuc_nang", 3: "nha_cung_cap", 4: "version", 5: "inter_conn", 6: "ghi_chu"}),
        (20, "K_van_ban", {0: "idx", 1: "loai", 2: "so", 3: "ghi_chu"}),
        (26, "R1_dao_tao", {0: "idx", 1: "hinh_thuc", 2: "to_chuc", 3: "thoi_gian", 4: "so_cb", 5: "chung_chi"}),
        (28, "S1_kiem_tra", {0: "idx", 1: "loai_kt", 2: "thuc_hien", 3: "thoi_gian", 4: "ket_qua"}),
        (31, "T2_port_mapping", {0: "sw_name", 1: "port_count", 2: "port_desc", 3: "ghi_chu"}),
        (33, "T5_vi_tri", {0: "ten_thiet_bi", 1: "tang", 2: "phong", 3: "ghi_chu"})
    ]

    for idx, name, mapping in tables_to_fix:
        clean_and_inject_triple_row(doc.tables[idx], name, mapping)

    # General replacements
    for p in doc.paragraphs:
        if "xã Lý Nhân, tỉnh Ninh Bình" in p.text:
            p.text = p.text.replace("xã Lý Nhân, tỉnh Ninh Bình", "{{ ten_don_vi }}")
        if "Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân" in p.text:
            p.text = p.text.replace("Hệ thống thông tin phục vụ hoạt động của UBND xã Lý Nhân", "{{ he_thong_thong_tin }}")
        if "thôn 5, xã Lý Nhân, tỉnh Ninh Bình" in p.text:
            p.text = p.text.replace("thôn 5, xã Lý Nhân, tỉnh Ninh Bình", "{{ dia_chi }}")

    doc.save(output_path)
    print(f"Triple Row Template successfully created: {output_path}")

if __name__ == "__main__":
    main()
