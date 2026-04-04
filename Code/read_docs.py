import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

os.chdir(r'd:\ToolKS\Code')
from docx import Document

# ===== FILE 1 =====
print("=" * 80)
print("FILE 1: HSDXCDAT LY NHAN")
print("=" * 80)

doc = Document('HSDXCDAT LY NHAN.docx.docx')

print('\n=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f'[P{i}] {p.style.name} | {t[:200]}')

print('\n=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'\n--- Table {ti} ({len(table.rows)}r x {len(table.columns)}c) ---')
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace('\n',' ')[:50] for c in row.cells]
        print(f'  R{ri}: {cells}')

print('\n=== IMAGES ===')
count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        count += 1
        print(f'  Img: {rel.target_ref}')
print(f'Total: {count}')

# ===== FILE 2 =====
print("\n" + "=" * 80)
print("FILE 2: BAO CAO HSDX")
print("=" * 80)

doc2 = Document('BAO CAO HSDX.docx.docx')

print('\n=== PARAGRAPHS ===')
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t:
        print(f'[P{i}] {p.style.name} | {t[:250]}')

print('\n=== TABLES ===')
for ti, table in enumerate(doc2.tables):
    print(f'\n--- Table {ti} ({len(table.rows)}r x {len(table.columns)}c) ---')
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace('\n',' ')[:60] for c in row.cells]
        print(f'  R{ri}: {cells}')
