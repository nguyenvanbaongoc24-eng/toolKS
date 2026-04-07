from docx import Document
import os

phieu = r"d:\ToolKS\Code\Phieu_Khao_Sat_ATTT_Mau_1.docx.docx"
doc = Document(phieu)

res = []
res.append(f"Total tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    res.append(f"\n--- Table {i} ---")
    if len(table.rows) > 0:
        first_row = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
        # remove duplicate columns
        seen = []
        unique_row = []
        for x in first_row:
            if x not in seen:
                seen.append(x)
                unique_row.append(x)
        res.append(f"Header: {unique_row}")
        
        if len(table.rows) > 1:
            second_row = [c.text.strip()[:30].replace('\n', ' ') for c in table.rows[1].cells]
            seen2 = []
            unique_row2 = []
            for x in second_row:
                if x not in seen2:
                    seen2.append(x)
                    unique_row2.append(x)
            res.append(f"Data: {unique_row2}")

with open(r"d:\ToolKS\Code\tables_info.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(res))
