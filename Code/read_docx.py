import docx
import io
import sys

def read_docx_to_file(path, out_path):
    try:
        doc = docx.Document(path)
        with io.open(out_path, "w", encoding="utf-8") as f:
            f.write("====== PARAGRAPHS ======\n")
            for p in doc.paragraphs:
                if p.text.strip():
                    f.write("P: " + p.text + "\n")
                    
            f.write("\n====== TABLES ======\n")
            for i, table in enumerate(doc.tables):
                f.write(f"\n--- Table {i+1} ---\n")
                for row in table.rows:
                    cells = [c.text.replace('\n', ' ').strip() for c in row.cells]
                    f.write(" | ".join(cells) + "\n")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        read_docx_to_file(sys.argv[1], sys.argv[2])
    else:
        read_docx_to_file('d:/ToolKS/Code/Phieu_Khao_Sat_ATTT_Mau_1.docx.docx', 'd:/ToolKS/Code/docx_output.txt')
