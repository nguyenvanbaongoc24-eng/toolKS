import docx
import os

def inject_checkboxes_full(out_dir):
    template_path = 'd:/ToolKS/Code/Phieu_Khao_Sat_ATTT_Mau_1.docx.docx'
    out_path = os.path.join(out_dir, 'phieu_ks_template_v2.docx') # Overwrite v2
    doc = docx.Document(template_path)
    
    # helper to replace patterns
    def r(old, new):
        for p in doc.paragraphs:
            if old in p.text:
                p.text = p.text.replace(old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old in p.text:
                            p.text = p.text.replace(old, new)

    # L1: Kiểm soát vật lý
    r("☐ Có khóa cửa (chìa khóa thường)", "{{ l1_phys_key }} Có khóa cửa")
    r("☐ Có khóa cửa + camera giám sát", "{{ l1_phys_cam }} Có khóa + camera")
    r("☐ Có thẻ từ / kiểm soát điện tử", "{{ l1_phys_card }} Có thẻ từ")
    r("☐ Không có kiểm soát riêng", "{{ l1_phys_none }} Không có kiểm soát")

    # L2: Logic
    r("☐ Có – Yêu cầu: độ dài tối thiểu", "{{ l2_pass_yes }} Có – Yêu cầu: độ dài")
    r("☐ Không có chính sách thống nhất", "{{ l2_pass_no }} Không có chính sách")
    r("_____ ký tự", "{{ l2_pass_len }} ký tự")
    r("_____ tháng", "{{ l2_pass_time }} tháng")

    # L3: Antivirus
    r("☐ Có – Tên phần mềm: _______________", "{{ l3_av_yes }} Có – Tên: {{ l3_av_name }}")
    r("☐ Không", "{{ l3_av_no }} Không")

    # L4: Backup
    r("☐ Có – Tần suất: Hàng ngày / Hàng tuần / Hàng tháng", "{{ l4_bak_yes }} Có - {{ l4_bak_freq }}")
    r("☐ Thực hiện thủ công khi nhớ", "{{ l4_bak_manual }} Thủ công")
    r("☐ Không sao lưu", "{{ l4_bak_none }} Không sao lưu")

    # L5: Log (Already in v2 but let's re-apply to be sure)
    r("☒ Có", "{{ l5_log_yes }} Có")
    r("☐ Không biết / Chưa kiểm tra", "{{ l5_log_unknown }} Không biết")
    r("☐ Không", "{{ l5_log_no }} Không")
    r("☐ < 3 tháng", "{{ l5_retention_3m }} < 3 tháng")
    r("☐ 3 – 6 tháng", "{{ l5_retention_6m }} 3-6 tháng")
    r("☐ > 6 tháng", "{{ l5_retention_gt6m }} > 6 tháng")
    r("☐ Không lưu", "{{ l5_retention_none }} Không lưu")

    # L7: Firewall
    r("☐ Tường lửa tích hợp trên Router", "{{ l7_type_spi }} Tường lửa tích hợp")
    r("☐ Tường lửa phần cứng chuyên dụng", "{{ l7_type_hardware }} Tường lửa phần cứng")
    r("☐ Tường lửa phần mềm trên máy chủ", "{{ l7_type_software }} Tường lửa phần mềm")

    # P: Mã hóa
    r("☐ HTTPS (có chứng chỉ SSL/TLS)", "{{ p1_https }} HTTPS")
    r("☐ HTTP (không mã hóa)", "{{ p1_http }} HTTP")
    r("☐ Cả hai", "{{ p1_both }} Cả hai")

    r("☐ Có – Loại VPN: ☐ SSL VPN  ☐ IPSec", "{{ p2_vpn_yes }} Có - {{ p2_vpn_type }}")
    r("☐ Không có VPN", "{{ p2_vpn_no }} Không có VPN")

    doc.save(out_path)
    print(f"Created FULL advanced template: {out_path}")

if __name__ == "__main__":
    out_dir = 'd:/ToolKS/Code/backend/templates'
    inject_checkboxes_full(out_dir)
