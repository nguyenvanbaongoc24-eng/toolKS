from docx import Document
import re, os

template_dir = r'd:\ToolKS\Code\backend\templates'
results = []

for fname in ['phieu_khao_sat_template.docx','hsdx_template.docx','bao_cao_template.docx']:
    path = os.path.join(template_dir, fname)
    doc = Document(path)
    tags = set()
    old_count = 0
    for p in doc.paragraphs:
        found = re.findall(r'\{\{[^}]+\}\}', p.text)
        tags.update(found)
        if 'Lý Nhân' in p.text:
            old_count += 1
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    found = re.findall(r'\{\{[^}]+\}\}', p.text)
                    tags.update(found)
                    if 'Lý Nhân' in p.text:
                        old_count += 1
    
    results.append(f"\n=== {fname} ===")
    results.append(f"Placeholders found: {len(tags)}")
    results.append(f"Old text remaining: {old_count}")
    for tag in sorted(tags):
        results.append(f"  {tag}")

with open(r'd:\ToolKS\Code\verify_result.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(results))

print("Done - check verify_result.txt")
